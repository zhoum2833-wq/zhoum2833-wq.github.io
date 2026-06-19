"""Generate Word document from all markdown files."""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import docx.oxml

docs_dir = os.path.join(os.path.dirname(__file__), 'docs')
output_path = os.path.join(os.path.dirname(__file__), '..', '电赛入门指南.docx')

CHAPTER_DIRS = [
    '00-intro', '01-hardware', '02-connection', '03-tools', '04-protocols',
    '05-internals', '06-rtos-linux', '07-programming', '08-ai', '09-pcb',
    '10-mechanical', '11-project-design', '12-modules'
]

CHAPTER_NAMES = {
    '00-intro': '第〇篇：认识单片机',
    '01-hardware': '第一篇：硬件基础',
    '02-connection': '第二篇：电脑如何连接单片机',
    '03-tools': '第三篇：开发工具与工程结构',
    '04-protocols': '第四篇：通信协议',
    '05-internals': '第五篇：单片机内部机制',
    '06-rtos-linux': '第六篇：RTOS 与 Linux',
    '07-programming': '第七篇：编程实践',
    '08-ai': '第八篇：AI 编程',
    '09-pcb': '第九篇：PCB 设计',
    '10-mechanical': '第十篇：机械结构',
    '11-project-design': '第十一篇：项目方案设计',
    '12-modules': '第十二篇：常见模块与实战',
}

CALLOUT_COLORS = {
    'tip': ('D5F5E3', '27AE60', '💡 提示'),
    'warning': ('FDEBD0', 'E67E22', '⚠️ 注意'),
    'danger': ('FADBD8', 'E74C3C', '🚫 警告'),
    'info': ('D6EAF8', '2980B9', 'ℹ️ 信息'),
}


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_code_block(doc, code_text):
    """Add a code block with gray background and Consolas font."""
    for line in code_text.split('\n'):
        p = doc.add_paragraph()
        p.style = doc.styles['Code']
        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')


def add_paragraph_with_inline(doc, text):
    """Add a paragraph with inline formatting (bold, code, links)."""
    p = doc.add_paragraph()
    p.style = doc.styles['Body Text']
    # Simple regex-based inline parsing
    remaining = text
    while remaining:
        # Bold
        m = re.match(r'\*\*(.+?)\*\*', remaining)
        if m:
            run = p.add_run(m.group(1))
            run.bold = True
            remaining = remaining[m.end():]
            continue
        # Inline code
        m = re.match(r'`([^`]+)`', remaining)
        if m:
            run = p.add_run(m.group(1))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8E8E8" w:val="clear"/>')
            remaining = remaining[m.end():]
            continue
        # Link
        m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', remaining)
        if m:
            run = p.add_run(m.group(1))
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.underline = True
            remaining = remaining[m.end():]
            continue
        # Plain text
        next_special = re.search(r'`|\*\*|\[', remaining)
        if not next_special:
            p.add_run(remaining)
            remaining = ''
        else:
            p.add_run(remaining[:next_special.start()])
            remaining = remaining[next_special.start():]


def add_table(doc, rows):
    """Add a table to the document."""
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True
                set_cell_shading(cell, 'D5E8F0')

    doc.add_paragraph()  # spacing after table


def add_callout(doc, callout_type, text):
    """Add a callout block with colored left border and background."""
    colors = CALLOUT_COLORS.get(callout_type, CALLOUT_COLORS['info'])
    bg, border_color, label = colors

    # Label
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    set_paragraph_border(p, border_color)
    set_paragraph_shading(p, bg)

    # Content
    for line in text.split('\n'):
        if line.strip():
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(10)
            set_paragraph_border(p, border_color)
            set_paragraph_shading(p, bg)

    doc.add_paragraph()  # spacing


def set_paragraph_border(paragraph, color):
    """Add left border to paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="8" w:space="4" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def set_paragraph_shading(paragraph, color):
    """Add background shading to paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    pPr.append(shd)


def parse_markdown(text):
    """Parse markdown text into blocks."""
    lines = text.split('\n')
    blocks = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith('```'):
            lang = line.strip()[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            blocks.append(('code', '\n'.join(code_lines)))
            continue

        # Table
        if '|' in line and line.strip().startswith('|'):
            table_rows = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].split('|') if c.strip()]
                table_rows.append(cells)
                i += 1
            # Filter separator rows
            filtered = [r for r in table_rows if not all(re.match(r'^[-:]+$', c) for c in r)]
            if filtered:
                blocks.append(('table', filtered))
            continue

        # Callout
        m = re.match(r'^:::\s*(tip|warning|danger|info)', line.strip())
        if m:
            callout_type = m.group(1)
            callout_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith(':::'):
                callout_lines.append(lines[i])
                i += 1
            i += 1  # skip closing :::
            blocks.append(('callout', callout_type, '\n'.join(callout_lines)))
            continue

        # Heading
        m = re.match(r'^(#{1,6})\s+(.+)', line)
        if m:
            level = len(m.group(1))
            blocks.append(('heading', min(level, 3), m.group(2)))
            i += 1
            continue

        # Empty line
        if line.strip() == '':
            i += 1
            continue

        # Normal paragraph
        para_lines = []
        while (i < len(lines) and lines[i].strip() != ''
               and not lines[i].strip().startswith('```')
               and not (lines[i].strip().startswith('|') and '|' in lines[i])
               and not re.match(r'^:::\s*(tip|warning|danger|info)', lines[i].strip())
               and not re.match(r'^#{1,6}\s', lines[i])):
            para_lines.append(lines[i])
            i += 1
        if para_lines:
            blocks.append(('paragraph', '\n'.join(para_lines)))

    return blocks


def setup_styles(doc):
    """Configure document styles."""
    # Normal style
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Body Text
    style = doc.styles['Body Text']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # Code style
    style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Consolas'
    style.font.size = Pt(9)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.1
    # Gray background
    pPr = style.element.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shd)

    # Heading styles
    for level, (size, name) in enumerate([(22, 'Heading 1'), (18, 'Heading 2'), (14, 'Heading 3')], 1):
        style = doc.styles[name]
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def add_cover_page(doc):
    """Add cover page."""
    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('电赛入门指南')
    run.font.size = Pt(36)
    run.bold = True
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('从零开始学嵌入式')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(85, 85, 85)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026 江苏省电子设计竞赛备赛教程')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(119, 119, 119)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('适用平台：MSPM0G3507 / STM32F103')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(153, 153, 153)

    doc.add_page_break()


def main():
    print("Creating Word document...")
    doc = Document()
    setup_styles(doc)

    # Add cover page
    add_cover_page(doc)

    # Add TOC placeholder
    p = doc.add_paragraph('目录', style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Add TOC field
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar)
    run = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run._r.append(instrText)
    run = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run._r.append(fldChar2)
    run = paragraph.add_run('（请在 Word 中右键此处，选择"更新域"以生成目录）')
    run.font.color.rgb = RGBColor(153, 153, 153)
    run.font.size = Pt(10)
    run = paragraph.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._r.append(fldChar3)

    doc.add_page_break()

    # Process all chapters
    total_files = 0
    for chapter_dir in CHAPTER_DIRS:
        dir_path = os.path.join(docs_dir, chapter_dir)
        if not os.path.isdir(dir_path):
            continue

        # Add chapter title
        chapter_name = CHAPTER_NAMES.get(chapter_dir, chapter_dir)
        doc.add_heading(chapter_name, level=1)

        # Process files
        files = sorted([f for f in os.listdir(dir_path) if f.endswith('.md')])
        for filename in files:
            filepath = os.path.join(dir_path, filename)
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()

            blocks = parse_markdown(text)
            for block in blocks:
                if block[0] == 'heading':
                    doc.add_heading(block[2], level=block[1])
                elif block[0] == 'paragraph':
                    add_paragraph_with_inline(doc, block[1])
                elif block[0] == 'code':
                    add_code_block(doc, block[1])
                    doc.add_paragraph()  # spacing
                elif block[0] == 'table':
                    add_table(doc, block[1])
                elif block[0] == 'callout':
                    add_callout(doc, block[1], block[2])

            total_files += 1
            if total_files % 10 == 0:
                print(f"  Processed {total_files} files...")

    # Add headers and footers to all sections
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

        # Header
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run('电赛入门指南')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(153, 153, 153)

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('第 ')
        run.font.size = Pt(9)
        # Page number field
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run._r.append(instrText)
        run = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run._r.append(fldChar2)
        run = p.add_run('1')
        run.font.size = Pt(9)
        run = p.add_run()
        fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._r.append(fldChar3)
        run = p.add_run(' 页')
        run.font.size = Pt(9)

    print(f"Saving to {output_path}...")
    doc.save(output_path)

    size_mb = os.path.getsize(output_path) / (1024 * 1024)
    print(f"Done! File: {output_path} ({size_mb:.1f} MB)")
    print(f"Total files processed: {total_files}")


if __name__ == '__main__':
    main()
