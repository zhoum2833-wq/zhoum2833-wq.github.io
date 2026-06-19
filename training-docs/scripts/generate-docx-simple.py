"""Generate Word document from 电赛入门指南.md."""
import os, re, gc
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC  = os.path.join(ROOT, '电赛入门指南.md')
OUT  = os.path.abspath(os.path.join(ROOT, 'output', '电赛入门指南.docx'))

CHAPTER_NAMES = {
    '00-intro': '第零篇：认识单片机',
    '01-hardware': '第一篇：硬件基础',
    '02-connection': '第二篇：电脑如何连接单片机',
    '03-tools': '第三篇：开发工具与工程结构',
    '04-protocols': '第四篇：通信协议',
    '05-internals': '第五篇：单片机内部机制',
    '06-rtos-linux': '第六篇：RTOS 与 Linux',
    '07-programming': '第七篇：编程实践',
    '08-ai': '第八篇：AI 编程',
    '09-pcb': '第九篇：PCB 设计',
    '10-mechanical': '第十篇：机械结构',
    '11-project-design': '第十一篇：项目方案设计',
    '12-modules': '第十二篇：常见模块与实战',
}


def add_para(doc, text, bold=False, size=11, font='Microsoft YaHei', code_font=False):
    p = doc.add_paragraph()
    if code_font:
        run = p.add_run(text if text.strip() else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>'))
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
    else:
        run = p.add_run(text)
        run.font.size = Pt(size)
        if bold:
            run.bold = True
    return p


def add_section_page(doc, title):
    doc.add_page_break()
    doc.add_heading(title, level=1)


def process_text(doc, body):
    """Process markdown text content and add to doc."""
    lines = body.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Code block
        if line.strip().startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                add_para(doc, lines[i].rstrip(), code_font=True)
                i += 1
            i += 1  # closing ```
            doc.add_paragraph()
            continue

        # Table
        if '|' in line and line.strip().startswith('|'):
            rows_raw = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].split('|') if c.strip()]
                rows_raw.append(cells)
                i += 1
            rows = [r for r in rows_raw if not all(re.match(r'^[-:]+$', c) for c in r)]
            if rows and len(rows) > 0:
                num_cols = len(rows[0])
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = 'Table Grid'
                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < num_cols:
                            cell = table.cell(ri, ci)
                            cell.text = ''
                            run = cell.paragraphs[0].add_run(cell_text)
                            run.font.size = Pt(9)
                            if ri == 0:
                                run.bold = True
                doc.add_paragraph()
            continue

        # Callout
        m = re.match(r'^:::\s*(tip|warning|danger|info)', line.strip())
        if m:
            ct = m.group(1)
            colors = {'tip': ('D5F5E3', '27AE60'), 'warning': ('FDEBD0', 'E67E22'),
                       'danger': ('FADBD8', 'E74C3C'), 'info': ('D6EAF8', '2980B9')}
            bg, bc = colors.get(ct, colors['info'])
            labels = {'tip': '💡 提示', 'warning': '⚠️ 注意', 'danger': '🚫 警告', 'info': 'ℹ️ 信息'}
            i += 1
            p = add_para(doc, labels.get(ct, '信息'), bold=True, size=10)
            pPr = p._p.get_or_add_pPr()
            pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>'))
            pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="8" w:space="4" w:color="{bc}"/></w:pBdr>'))
            while i < len(lines) and not lines[i].strip().startswith(':::'):
                if lines[i].strip():
                    p = add_para(doc, lines[i].strip(), size=10)
                    pPr = p._p.get_or_add_pPr()
                    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>'))
                    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="8" w:space="4" w:color="{bc}"/></w:pBdr>'))
                i += 1
            i += 1
            doc.add_paragraph()
            continue

        # Heading (skip # headings that are chapter titles, they're handled elsewhere)
        m = re.match(r'^(#{1,6})\s+(.+)', line)
        if m:
            level = min(len(m.group(1)), 3)
            doc.add_heading(m.group(2), level=level)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        clean = re.sub(r'`([^`]+)`', r'\1', clean)
        clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean)
        add_para(doc, clean)
        i += 1


def main():
    if not os.path.exists(SRC):
        print(f"ERROR: {SRC} not found")
        return

    with open(SRC, 'r', encoding='utf-8') as f:
        content = f.read()

    print("Creating Word document...")
    doc = Document()

    # Style setup
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.3
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Cover page
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('电赛入门指南')
    run.font.size = Pt(36)
    run.bold = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('从零开始学嵌入式').font.size = Pt(22)
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('2026 江苏省电子设计竞赛备赛教程')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_page_break()

    # TOC page
    doc.add_heading('目录', level=1)
    p = doc.add_paragraph('（在 Word 中右键此处 - 更新域 - 更新整个目录，即可自动生成完整目录）')
    p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
    doc.add_page_break()

    # Split by @split markers
    pattern = re.compile(r'<!-- @split: (.+?) -->\n', re.MULTILINE)
    parts = pattern.split(content)

    if len(parts) < 2:
        print("ERROR: No @split markers found")
        return

    count = 0
    last_chapter = ''
    for i in range(1, len(parts), 2):
        path = parts[i]
        body = parts[i + 1].strip()

        # path format: "00-intro/what-is-mcu.md"
        match = re.match(r'^([^/]+)/(.+)$', path)
        if not match:
            continue

        ch_dir = match.group(1)

        # Add chapter section page when chapter changes
        if ch_dir != last_chapter and ch_dir in CHAPTER_NAMES:
            add_section_page(doc, CHAPTER_NAMES[ch_dir])
            last_chapter = ch_dir

        # Remove trailing ---
        body = re.sub(r'\n---\s*$', '', body)

        process_text(doc, body)
        count += 1
        if count % 10 == 0:
            print(f"  {count} files...")
        if count % 20 == 0:
            gc.collect()

    # Page margins / header / footer
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hr = hp.add_run('电赛入门指南')
        hr.font.size = Pt(8)
        hr.font.color.rgb = RGBColor(150, 150, 150)
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.add_run('第 ').font.size = Pt(8)
        run = fp.add_run()
        run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        run = fp.add_run()
        run._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        run = fp.add_run()
        run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
        fp.add_run('1').font.size = Pt(8)
        run = fp.add_run()
        run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
        fp.add_run(' 页').font.size = Pt(8)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    print(f"Saving {OUT}...")
    doc.save(OUT)
    size = os.path.getsize(OUT) / 1024 / 1024
    print(f"Done! {OUT} ({size:.1f} MB), {count} articles processed.")


if __name__ == '__main__':
    main()
